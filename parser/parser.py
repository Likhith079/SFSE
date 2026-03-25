import os
from typing import Optional


def parse_file(path: str) -> str:
    """Dispatch to the correct parser based on file extension."""
    ext = os.path.splitext(path)[1].lower()

    if ext == ".txt":
        return _parse_txt(path)
    elif ext == ".pdf":
        return _parse_pdf(path)
    elif ext == ".docx":
        return _parse_docx(path)
    else:
        return ""


# ── Plain text ─────────────────────────────────────────────────────────────────

def _parse_txt(path: str) -> str:
    encodings = ["utf-8", "latin-1", "cp1252"]
    for enc in encodings:
        try:
            with open(path, "r", encoding=enc, errors="replace") as fh:
                return fh.read()
        except (OSError, UnicodeDecodeError):
            continue
    return ""


# ── PDF ────────────────────────────────────────────────────────────────────────

def _parse_pdf(path: str) -> str:
    try:
        import PyPDF2  # noqa: F401
    except ImportError:
        print("[parser] PyPDF2 not installed – skipping PDF:", path)
        return ""

    text_parts = []
    try:
        with open(path, "rb") as fh:
            reader = PyPDF2.PdfReader(fh)
            for page in reader.pages:
                try:
                    text_parts.append(page.extract_text() or "")
                except Exception:
                    pass
    except Exception as exc:
        print(f"[parser] PDF error {path}: {exc}")

    return "\n".join(text_parts)


# ── DOCX ───────────────────────────────────────────────────────────────────────

def _parse_docx(path: str) -> str:
    try:
        from docx import Document  # noqa: F401
    except ImportError:
        print("[parser] python-docx not installed – skipping DOCX:", path)
        return ""

    text_parts = []
    try:
        doc = Document(path)
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
    except Exception as exc:
        print(f"[parser] DOCX error {path}: {exc}")

    return "\n".join(text_parts)